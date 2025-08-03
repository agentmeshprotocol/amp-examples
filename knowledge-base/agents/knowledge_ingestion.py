"""
Knowledge Ingestion Agent

This agent handles document processing, text extraction, chunking,
embedding generation, and indexing for the knowledge base system.
"""

import asyncio
import hashlib
import logging
import mimetypes
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF for better PDF processing
import nltk
import numpy as np
import spacy
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..', '..', 'shared-lib'))
from amp_client import AMPClient
from amp_types import AMPMessage, MessageType

# Download required NLTK data
try:
    nltk.data.find('punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('stopwords')
except LookupError:
    nltk.download('stopwords')

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document processing and text extraction"""
    
    def __init__(self):
        self.nlp = spacy.load("en_core_web_sm")
        
    async def extract_text(self, file_path: str, mime_type: Optional[str] = None) -> Tuple[str, Dict[str, Any]]:
        """Extract text from various document formats"""
        if not mime_type:
            mime_type, _ = mimetypes.guess_type(file_path)
            
        metadata = {
            'file_path': file_path,
            'mime_type': mime_type,
            'file_size': os.path.getsize(file_path),
            'processed_at': datetime.utcnow().isoformat()
        }
        
        try:
            if mime_type == 'application/pdf':
                text, pdf_metadata = await self._extract_pdf_text(file_path)
                metadata.update(pdf_metadata)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                text, docx_metadata = await self._extract_docx_text(file_path)
                metadata.update(docx_metadata)
            elif mime_type in ['text/html', 'application/xhtml+xml']:
                text, html_metadata = await self._extract_html_text(file_path)
                metadata.update(html_metadata)
            elif mime_type in ['text/plain', 'text/markdown']:
                text, txt_metadata = await self._extract_text_file(file_path)
                metadata.update(txt_metadata)
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
                
            # Clean and normalize text
            text = self._clean_text(text)
            metadata['text_length'] = len(text)
            metadata['word_count'] = len(text.split())
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF files"""
        doc = fitz.open(file_path)
        text_parts = []
        metadata = {
            'page_count': len(doc),
            'pdf_metadata': doc.metadata
        }
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_parts.append(page.get_text())
            
        doc.close()
        return '\n'.join(text_parts), metadata
    
    async def _extract_docx_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX files"""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
            
        metadata = {
            'paragraph_count': len(doc.paragraphs)
        }
        
        return '\n'.join(text_parts), metadata
    
    async def _extract_html_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from HTML files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        soup = BeautifulSoup(content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        text = soup.get_text()
        
        metadata = {
            'title': soup.title.string if soup.title else None,
            'meta_description': None
        }
        
        # Extract meta description
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc:
            metadata['meta_description'] = meta_desc.get('content')
            
        return text, metadata
    
    async def _extract_text_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            
        metadata = {
            'encoding': 'utf-8',
            'line_count': len(text.split('\n'))
        }
        
        return text, metadata
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
                
        return '\n'.join(cleaned_lines)


class TextChunker:
    """Handles intelligent text chunking for embedding generation"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.nlp = spacy.load("en_core_web_sm")
        
    async def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split text into semantically meaningful chunks"""
        doc = self.nlp(text)
        sentences = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        
        chunks = []
        current_chunk = ""
        current_length = 0
        chunk_id = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # If adding this sentence would exceed chunk size, finalize current chunk
            if current_length + sentence_length > self.chunk_size and current_chunk:
                chunk_data = await self._create_chunk(
                    current_chunk.strip(), 
                    chunk_id, 
                    metadata
                )
                chunks.append(chunk_data)
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence
                current_length = len(current_chunk)
                chunk_id += 1
            else:
                current_chunk += " " + sentence
                current_length += sentence_length + 1
        
        # Add final chunk
        if current_chunk.strip():
            chunk_data = await self._create_chunk(
                current_chunk.strip(), 
                chunk_id, 
                metadata
            )
            chunks.append(chunk_data)
            
        return chunks
    
    async def _create_chunk(self, text: str, chunk_id: int, source_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create a chunk with metadata"""
        chunk_hash = hashlib.md5(text.encode()).hexdigest()
        
        # Extract entities and keywords
        doc = self.nlp(text)
        entities = [(ent.text, ent.label_) for ent in doc.ents]
        keywords = [token.lemma_.lower() for token in doc 
                   if not token.is_stop and not token.is_punct and token.is_alpha]
        
        return {
            'id': f"{source_metadata.get('file_path', 'unknown')}_{chunk_id}",
            'text': text,
            'chunk_id': chunk_id,
            'hash': chunk_hash,
            'length': len(text),
            'word_count': len(text.split()),
            'entities': entities,
            'keywords': list(set(keywords)),
            'source_metadata': source_metadata,
            'created_at': datetime.utcnow().isoformat()
        }
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text for chunk continuity"""
        words = text.split()
        overlap_words = words[-self.chunk_overlap//10:] if len(words) > self.chunk_overlap//10 else words
        return " ".join(overlap_words)


class EmbeddingGenerator:
    """Generates embeddings for text chunks"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model = SentenceTransformer(model_name)
        self.model_name = model_name
        
    async def generate_embeddings(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Generate embeddings for text chunks"""
        texts = [chunk['text'] for chunk in chunks]
        
        # Generate embeddings in batches for efficiency
        batch_size = 32
        all_embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i + batch_size]
            batch_embeddings = self.model.encode(batch_texts, convert_to_tensor=False)
            all_embeddings.extend(batch_embeddings)
        
        # Add embeddings to chunks
        for chunk, embedding in zip(chunks, all_embeddings):
            chunk['embedding'] = embedding.tolist()
            chunk['embedding_model'] = self.model_name
            chunk['embedding_dimension'] = len(embedding)
            
        return chunks


class KnowledgeIngestionAgent:
    """Main knowledge ingestion agent that orchestrates document processing"""
    
    def __init__(self, agent_id: str = "knowledge-ingestion-agent"):
        self.agent_id = agent_id
        self.client = AMPClient(agent_id)
        
        # Initialize processors
        self.document_processor = DocumentProcessor()
        self.text_chunker = TextChunker()
        self.embedding_generator = EmbeddingGenerator()
        
        # Configuration
        self.supported_formats = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/html',
            'application/xhtml+xml',
            'text/plain',
            'text/markdown'
        ]
        
        logger.info(f"Knowledge Ingestion Agent {agent_id} initialized")
    
    async def start(self, host: str = "localhost", port: int = 8000):
        """Start the agent and register capabilities"""
        await self.client.connect(f"ws://{host}:{port}/ws")
        
        # Register capabilities
        capabilities = [
            {
                "name": "document-ingestion",
                "description": "Process and ingest documents into knowledge base",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "file_path": {"type": "string"},
                        "metadata": {"type": "object"},
                        "processing_options": {"type": "object"}
                    },
                    "required": ["file_path"]
                },
                "output_schema": {
                    "type": "object",
                    "properties": {
                        "document_id": {"type": "string"},
                        "chunks": {"type": "array"},
                        "metadata": {"type": "object"},
                        "status": {"type": "string"}
                    }
                }
            },
            {
                "name": "batch-ingestion",
                "description": "Process multiple documents in batch",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "file_paths": {"type": "array", "items": {"type": "string"}},
                        "processing_options": {"type": "object"}
                    },
                    "required": ["file_paths"]
                }
            },
            {
                "name": "text-processing",
                "description": "Process raw text without file input",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "text": {"type": "string"},
                        "metadata": {"type": "object"}
                    },
                    "required": ["text"]
                }
            }
        ]
        
        for capability in capabilities:
            await self.client.register_capability(capability)
        
        # Start message handling
        await self.client.start_message_handler(self._handle_message)
        logger.info(f"Knowledge Ingestion Agent started on {host}:{port}")
    
    async def _handle_message(self, message: AMPMessage):
        """Handle incoming AMP messages"""
        try:
            capability = message.message.destination.capability
            payload = message.message.payload
            
            if capability == "document-ingestion":
                result = await self._handle_document_ingestion(payload)
            elif capability == "batch-ingestion":
                result = await self._handle_batch_ingestion(payload)
            elif capability == "text-processing":
                result = await self._handle_text_processing(payload)
            else:
                raise ValueError(f"Unknown capability: {capability}")
            
            # Send response
            await self.client.send_response(
                message.message.id,
                message.message.source.agent_id,
                result
            )
            
        except Exception as e:
            logger.error(f"Error handling message: {e}")
            await self.client.send_error(
                message.message.id,
                message.message.source.agent_id,
                str(e),
                "PROCESSING_ERROR"
            )
    
    async def _handle_document_ingestion(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        """Handle single document ingestion"""
        file_path = payload['file_path']
        additional_metadata = payload.get('metadata', {})
        processing_options = payload.get('processing_options', {})
        
        # Check if file exists and is supported
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {mime_type}")
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        
        # Extract text
        text, metadata = await self.document_processor.extract_text(file_path, mime_type)
        metadata.update(additional_metadata)
        metadata['document_id'] = document_id
        
        # Chunk text
        chunk_size = processing_options.get('chunk_size', 1000)
        chunk_overlap = processing_options.get('chunk_overlap', 200)
        
        chunker = TextChunker(chunk_size, chunk_overlap)
        chunks = await chunker.chunk_text(text, metadata)
        
        # Generate embeddings
        if processing_options.get('generate_embeddings', True):
            chunks = await self.embedding_generator.generate_embeddings(chunks)
        
        # Send chunks to semantic search agent for indexing
        await self._send_for_indexing(chunks)
        
        # Send to knowledge graph agent for entity extraction
        await self._send_for_graph_processing(chunks)
        
        return {
            "document_id": document_id,
            "chunks": chunks,
            "metadata": metadata,
            "status": "processed",
            "chunk_count": len(chunks)
        }
    
    async def _handle_batch_ingestion(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        """Handle batch document ingestion"""
        file_paths = payload['file_paths']
        processing_options = payload.get('processing_options', {})
        
        results = []
        errors = []
        
        for file_path in file_paths:
            try:
                result = await self._handle_document_ingestion({
                    'file_path': file_path,
                    'processing_options': processing_options
                })
                results.append(result)
            except Exception as e:
                errors.append({
                    'file_path': file_path,
                    'error': str(e)
                })
        
        return {
            "status": "completed",
            "processed_count": len(results),
            "error_count": len(errors),
            "results": results,
            "errors": errors
        }
    
    async def _handle_text_processing(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        """Handle raw text processing"""
        text = payload['text']
        metadata = payload.get('metadata', {})
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        metadata['document_id'] = document_id
        metadata['source_type'] = 'raw_text'
        metadata['processed_at'] = datetime.utcnow().isoformat()
        
        # Chunk text
        chunks = await self.text_chunker.chunk_text(text, metadata)
        
        # Generate embeddings
        chunks = await self.embedding_generator.generate_embeddings(chunks)
        
        # Send for indexing and graph processing
        await self._send_for_indexing(chunks)
        await self._send_for_graph_processing(chunks)
        
        return {
            "document_id": document_id,
            "chunks": chunks,
            "metadata": metadata,
            "status": "processed",
            "chunk_count": len(chunks)
        }
    
    async def _send_for_indexing(self, chunks: List[Dict[str, Any]]):
        """Send chunks to semantic search agent for indexing"""
        try:
            await self.client.send_request(
                "semantic-search-agent",
                "index-chunks",
                {"chunks": chunks}
            )
        except Exception as e:
            logger.warning(f"Failed to send chunks for indexing: {e}")
    
    async def _send_for_graph_processing(self, chunks: List[Dict[str, Any]]):
        """Send chunks to knowledge graph agent for processing"""
        try:
            await self.client.send_request(
                "knowledge-graph-agent",
                "extract-entities",
                {"chunks": chunks}
            )
        except Exception as e:
            logger.warning(f"Failed to send chunks for graph processing: {e}")
    
    async def stop(self):
        """Stop the agent"""
        await self.client.disconnect()
        logger.info("Knowledge Ingestion Agent stopped")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Knowledge Ingestion Agent")
    parser.add_argument("--host", default="localhost", help="Host to connect to")
    parser.add_argument("--port", type=int, default=8000, help="Port to connect to")
    parser.add_argument("--agent-id", default="knowledge-ingestion-agent", help="Agent ID")
    
    args = parser.parse_args()
    
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    agent = KnowledgeIngestionAgent(args.agent_id)
    
    try:
        asyncio.run(agent.start(args.host, args.port))
    except KeyboardInterrupt:
        asyncio.run(agent.stop())